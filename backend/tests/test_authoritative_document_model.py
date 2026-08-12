from io import BytesIO
from pathlib import Path

from docx import Document

from app.engine.analyzer import ThreatAnalyzer
from app.engine.parser import ArchitectureParser
from app.services import document_ingestion


ROOT = Path(__file__).resolve().parents[2]
COMPLEX_SCENARIO = ROOT / "test-documents" / "Complex_Healthcare_AI_Threat_Model_Scenario.docx"


def _extracted_scenario():
    text, metadata = document_ingestion._extract_docx_text(COMPLEX_SCENARIO.read_bytes())
    description = (
        f"Document: {COMPLEX_SCENARIO.name}\n"
        "Type: docx\nRole: source_design\nContent:\n"
        f"{text}"
    )
    return description, metadata


def test_docx_extraction_preserves_table_order_and_structure():
    description, metadata = _extracted_scenario()

    assert metadata["extraction_quality"] == "structured_text_complete"
    assert metadata["tables"] == "15"
    assert "[Table 5]" in description
    assert "C25 | Delivery platform" in description
    assert "F20 | Support engineer -> support portal -> tenant data" in description
    assert "K30 | Privacy / deletion" in description
    assert description.index("3. Architecture inventory") < description.index("[Table 5]")
    assert description.index("8. Known weaknesses") < description.index("K1 | Authorization")


def test_docx_fallback_preserves_tables_without_python_docx(monkeypatch):
    document = Document()
    document.add_paragraph("3. Architecture inventory")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "ID"
    table.rows[0].cells[1].text = "Component"
    table.rows[0].cells[2].text = "Technology"
    table.rows[1].cells[0].text = "C1"
    table.rows[1].cells[1].text = "Core API"
    table.rows[1].cells[2].text = "Node.js"
    payload = BytesIO()
    document.save(payload)

    monkeypatch.setattr(document_ingestion, "Document", None)
    text, metadata = document_ingestion._extract_docx_text(payload.getvalue())

    assert metadata["extraction_quality"] == "structured_text_fallback"
    assert metadata["tables"] == "1"
    assert text.index("3. Architecture inventory") < text.index("[Table 1]")
    assert "Row 2: C1 | Core API | Node.js" in text


def test_authoritative_tables_replace_heuristic_topology():
    description, _ = _extracted_scenario()
    architecture = ArchitectureParser().parse(description)

    technical = [component for component in architecture.components if component.id.startswith("c")]
    external = [component for component in architecture.components if component.id.startswith("ext_")]
    flow_ids = {flow.properties.get("source_record_id") for flow in architecture.flows}
    component_ids = {component.id for component in architecture.components}

    assert len(technical) == 25
    assert len(external) == 3
    assert len(architecture.flows) == 20
    assert flow_ids == {f"F{index}" for index in range(1, 21)}
    assert all(not flow.assumed for flow in architecture.flows)
    assert len(architecture.trust_boundaries) == 9
    assert len(architecture.assets) == 12
    assert len(architecture.metadata["actors"]) == 10
    assert len(architecture.metadata["known_issues"]) == 30
    assert "ups_external" not in component_ids
    assert "vault" not in component_ids
    assert next(item for item in technical if item.id == "c16").trust_level == "restricted"
    assert next(item for item in technical if item.id == "c15").properties["public_access"] is False


def test_complex_scenario_findings_are_complete_and_grounded():
    description, metadata = _extracted_scenario()
    result = ThreatAnalyzer().analyze_from_text(
        description,
        project_name="Complex document regression",
        use_local_slm=False,
        analysis_mode="standard",
        domain_profile="healthcare",
        source_documents=[{
            "filename": COMPLEX_SCENARIO.name,
            "type": "docx",
            "role": "source_design",
            **metadata,
        }],
    )

    confirmed = [threat for threat in result.threats if threat.tier == "Confirmed"]
    potential = [threat for threat in result.threats if threat.tier == "Potential"]
    assert len(confirmed) == 30
    assert len(potential) == 6
    assert all(any(detail.get("source_ref", "").startswith("K") for detail in threat.evidence_details)
               for threat in confirmed)
    assert {threat.stride_category for threat in result.threats} == {
        "Spoofing", "Tampering", "Repudiation", "Information Disclosure",
        "Denial of Service", "Elevation of Privilege",
    }
    assert not any(threat.id.startswith("KB-K8S-005") for threat in result.threats)
    assert all(threat.id.startswith("STRIDE-") for threat in potential)
    assert not any("public access" in threat.title.lower() and threat.affected_component == "c15"
                   for threat in result.threats)

    known_ids = {threat.id.rsplit("-K", 1)[0] for threat in confirmed}
    assert {
        "API-BOLA-TENANT-CONTROL-001",
        "AUTH-SESSION-REVOCATION-001",
        "FHIR-PARTNER-SPOOFING-001",
        "WEB-SQL-INJECTION-ORDER-001",
        "WEB-STORED-XSS-001",
        "WEB-SSRF-URL-FETCH-001",
        "AI-RAG-TENANT-ISOLATION-001",
        "AI-INDIRECT-PROMPT-INJECTION-001",
        "MCP-DELEGATED-AUTHORIZATION-001",
        "PAYMENT-IDEMPOTENCY-001",
        "SUPPLY-CHAIN-GITHUB-OIDC-001",
        "DATA-DELETION-PROPAGATION-001",
    }.issubset(known_ids)

    assert result.architecture_validation["counts"] == {
        "components": 28,
        "explicit_components": 28,
        "flows": 20,
        "explicit_flows": 20,
        "inferred_flows": 0,
        "actors": 10,
        "identities": 10,
    }
    assert "React patient web app" in result.mermaid_diagram
    assert "Public API edge" in result.mermaid_diagram
    assert "Core API" in result.mermaid_diagram
    assert "Transactional database" in result.mermaid_diagram
    assert "AI orchestrator" not in result.mermaid_diagram
