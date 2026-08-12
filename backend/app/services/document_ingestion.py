import io
import csv
import json
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, List, Tuple

try:
    import yaml
except Exception:  # pragma: no cover
    yaml = None

from fastapi import UploadFile

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency at runtime
    PdfReader = None

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
except Exception:  # pragma: no cover - optional dependency at runtime
    Document = None
    Table = Paragraph = CT_Tbl = CT_P = None


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".adoc",
    ".csv",
    ".log",
    ".json",
    ".yaml",
    ".yml",
}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx"}
MAX_DOCUMENT_BYTES = 8 * 1024 * 1024


def _read_text_bytes(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode the uploaded text document.")


def _extract_pdf_text(raw_bytes: bytes) -> Tuple[str, Dict[str, str]]:
    if PdfReader is None:
        raise ValueError("PDF support is not installed on the backend.")

    reader = PdfReader(io.BytesIO(raw_bytes))
    pages = []
    image_only_pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {index}]\n{text}")
        else:
            image_only_pages.append(str(index))

    if not pages:
        raise ValueError("The uploaded PDF did not contain extractable text.")

    details = {
        "pages": str(len(reader.pages)),
        "image_only_pages": ",".join(image_only_pages),
        "extraction_quality": "partial" if image_only_pages else "text_complete",
        "warning": (
            f"Pages {', '.join(image_only_pages)} contained no extractable text; OCR or diagram review is required."
            if image_only_pages else ""
        ),
    }
    return "\n\n".join(pages), details


def _extract_docx_text(raw_bytes: bytes) -> Tuple[str, Dict[str, str]]:
    if Document is not None:
        document = Document(io.BytesIO(raw_bytes))
        blocks = []
        paragraph_count = 0
        table_count = 0
        # Iterating document.paragraphs followed by document.tables destroys
        # section context. Known-issue and architecture tables must remain next
        # to their headings so the architecture parser can treat them as
        # authoritative records.
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, document)
                text = paragraph.text.strip()
                if text:
                    paragraph_count += 1
                    blocks.append(f"[Paragraph {paragraph_count}] {text}")
            elif isinstance(child, CT_Tbl):
                table_count += 1
                table = Table(child, document)
                blocks.append(f"[Table {table_count}]")
                for row_index, row in enumerate(table.rows, 1):
                    values = [
                        re.sub(r"\s+", " ", cell.text).strip().replace("|", "&#124;")
                        for cell in row.cells
                    ]
                    blocks.append(f"Row {row_index}: " + " | ".join(values))

        if blocks:
            image_count = len(document.inline_shapes)
            return "\n".join(blocks), {
                "paragraphs": str(paragraph_count),
                "tables": str(table_count),
                "embedded_images": str(image_count),
                "extraction_quality": "partial" if image_count else "structured_text_complete",
                "warning": "Embedded images require separate diagram review." if image_count else "",
            }

    # Fallback parser for environments where python-docx is missing.
    try:
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as archive:
            xml_payload = archive.read("word/document.xml")
        root = ET.fromstring(xml_payload)
        body = root.find("w:body", namespace)
        if body is None:
            raise ValueError("DOCX body is missing.")

        blocks = []
        paragraph_count = 0
        table_count = 0
        for child in list(body):
            local_name = child.tag.rsplit("}", 1)[-1]
            if local_name == "p":
                texts = [node.text for node in child.findall(".//w:t", namespace) if node.text]
                text = "".join(texts).strip()
                if text:
                    paragraph_count += 1
                    blocks.append(f"[Paragraph {paragraph_count}] {text}")
            elif local_name == "tbl":
                table_count += 1
                blocks.append(f"[Table {table_count}]")
                for row_index, row in enumerate(child.findall("./w:tr", namespace), 1):
                    values = []
                    for cell in row.findall("./w:tc", namespace):
                        paragraphs = []
                        for paragraph in cell.findall(".//w:p", namespace):
                            texts = [node.text for node in paragraph.findall(".//w:t", namespace) if node.text]
                            if texts:
                                paragraphs.append("".join(texts))
                        value = re.sub(r"\s+", " ", " ".join(paragraphs)).strip().replace("|", "&#124;")
                        values.append(value)
                    blocks.append(f"Row {row_index}: " + " | ".join(values))
    except Exception as exc:
        raise ValueError(f"DOCX support is not installed on the backend and fallback parsing failed: {exc}") from exc

    if not blocks:
        raise ValueError("The uploaded DOCX did not contain extractable text.")

    return "\n".join(blocks), {
        "paragraphs": str(paragraph_count), "tables": str(table_count), "embedded_images": "unknown",
        "extraction_quality": "structured_text_fallback",
        "warning": "DOCX images require separate review; paragraph and table order was preserved.",
    }


def _extract_structured_text(raw_bytes: bytes, extension: str) -> Tuple[str, Dict[str, str]]:
    raw_text = _read_text_bytes(raw_bytes).strip()
    if extension == ".json":
        parsed = json.loads(raw_text)
        flattened = _flatten_structure(parsed)
        return json.dumps(parsed, indent=2, ensure_ascii=True) + "\n\n[Structured paths]\n" + "\n".join(flattened), {
            "extraction_quality": "structured_complete", "structured_format": "json", "warning": "",
        }
    if extension in {".yaml", ".yml"} and yaml is not None:
        parsed = yaml.safe_load(raw_text)
        flattened = _flatten_structure(parsed)
        return raw_text + "\n\n[Structured paths]\n" + "\n".join(flattened), {
            "extraction_quality": "structured_complete", "structured_format": "yaml", "warning": "",
        }
    if extension == ".csv":
        rows = list(csv.reader(io.StringIO(raw_text)))
        rendered = [f"[CSV row {index}] " + " | ".join(row) for index, row in enumerate(rows, 1)]
        return "\n".join(rendered), {
            "extraction_quality": "structured_complete", "structured_format": "csv", "rows": str(len(rows)), "warning": "",
        }
    return raw_text, {"extraction_quality": "text_complete", "warning": ""}


def _flatten_structure(value, path: str = "$") -> List[str]:
    lines = []
    if isinstance(value, dict):
        for key, child in value.items():
            lines.extend(_flatten_structure(child, f"{path}.{key}"))
    elif isinstance(value, list):
        for index, child in enumerate(value):
            lines.extend(_flatten_structure(child, f"{path}[{index}]"))
    else:
        lines.append(f"{path} = {value!r}")
    return lines


def _extract_text_from_bytes(filename: str, raw_bytes: bytes) -> Tuple[str, str, Dict[str, str]]:
    _, extension = os.path.splitext((filename or "").lower())
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{extension or 'unknown'}'. "
            "Supported formats: txt, md, rst, csv, json, yaml, yml, pdf, docx."
        )

    if len(raw_bytes) > MAX_DOCUMENT_BYTES:
        raise ValueError(f"File '{filename}' exceeds the 8 MB upload limit.")

    if extension in TEXT_EXTENSIONS:
        text, details = _extract_structured_text(raw_bytes, extension)
        return text, extension, details
    if extension == ".pdf":
        text, details = _extract_pdf_text(raw_bytes)
        return text.strip(), extension, details
    if extension == ".docx":
        text, details = _extract_docx_text(raw_bytes)
        return text.strip(), extension, details

    raise ValueError(f"Unsupported file type '{extension}'.")


async def extract_documents(files: List[UploadFile]) -> Tuple[str, List[Dict[str, str]]]:
    if not files:
        raise ValueError("At least one design document must be uploaded.")

    extracted_sections: List[str] = []
    metadata: List[Dict[str, str]] = []

    for file in files:
        filename = file.filename or "uploaded-document"
        raw_bytes = await file.read()
        extracted_text, extension, extraction_details = _extract_text_from_bytes(filename, raw_bytes)
        if not extracted_text:
            raise ValueError(f"File '{filename}' did not contain usable text.")

        lowered = extracted_text.lower()
        role = "source_design"
        if (
            "threat modeling report" in lowered
            or "security score" in lowered and "confirmed risks" in lowered
            or "top 3 things to fix first" in lowered
        ):
            role = "reference_report"

        extracted_sections.append(
            f"Document: {filename}\n"
            f"Type: {extension.lstrip('.')}\n"
            f"Role: {role}\n"
            f"Content:\n{extracted_text}"
        )
        metadata.append(
            {
                "filename": filename,
                "type": extension.lstrip("."),
                "role": role,
                "characters": str(len(extracted_text)),
                **extraction_details,
            }
        )

    return "\n\n---\n\n".join(extracted_sections), metadata
